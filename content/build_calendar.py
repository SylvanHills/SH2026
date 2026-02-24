#!/usr/bin/env python3
"""Build Sylvan Hills content calendar as .xlsx and .docx files."""

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import CellIsRule
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Palette ────────────────────────────────────────────────────────────────
DEEP_NAVY   = "1B2A4A"
WARM_TAUPE  = "C4AA8A"
MID_TAUPE   = "9E8A6E"
SOFT_NEUTRAL= "F5F0E8"
CREAM       = "EDE8DF"
LIGHT_GREY  = "D8D8D8"
SAGE_GREEN  = "8FAF8F"
SOFT_NAVY   = "4A6080"
WHITE       = "FFFFFF"
DARK_TEXT   = "2C2C2C"
TAUPE_BORDER= "C9B99A"
OFF_WHITE   = "FAFAF7"

# ─── Shared cell helpers ─────────────────────────────────────────────────────
THIN_SIDE   = Side(style="thin", color=TAUPE_BORDER)
THIN_BORDER = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

HEADER_FILL     = PatternFill("solid", fgColor=DEEP_NAVY)
WEEK_FILL       = PatternFill("solid", fgColor=MID_TAUPE)
ALT_FILL        = PatternFill("solid", fgColor=CREAM)
NORMAL_FILL     = PatternFill("solid", fgColor=SOFT_NEUTRAL)
OFF_WHITE_FILL  = PatternFill("solid", fgColor=OFF_WHITE)

HEADER_FONT     = Font(name="Calibri", bold=True,   color=WHITE,     size=11)
WEEK_FONT       = Font(name="Calibri", bold=True,   color=WHITE,     size=10)
BODY_FONT       = Font(name="Calibri",              color=DARK_TEXT, size=10)
TITLE_FONT      = Font(name="Calibri", bold=True,   color=DEEP_NAVY, size=14)
SUBTITLE_FONT   = Font(name="Calibri", italic=True, color=MID_TAUPE, size=10)
PROMPT_FONT     = Font(name="Calibri", bold=True,   color=DEEP_NAVY, size=10)

def hdr(cell):
    cell.fill, cell.font = HEADER_FILL, HEADER_FONT
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = THIN_BORDER

def body(cell, alt=False):
    cell.fill = ALT_FILL if alt else NORMAL_FILL
    cell.font = BODY_FONT
    cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    cell.border = THIN_BORDER

def section_hdr(ws, row, text, end_col="D"):
    ws.merge_cells(f"A{row}:{end_col}{row}")
    c = ws.cell(row=row, column=1, value=text)
    c.fill = HEADER_FILL
    c.font = Font(name="Calibri", bold=True, color=WHITE, size=11)
    c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    c.border = THIN_BORDER
    ws.row_dimensions[row].height = 26

def prompt_row(ws, row, label, end_col="D", alt=False):
    c = ws.cell(row=row, column=1, value=label)
    c.fill = ALT_FILL if alt else NORMAL_FILL
    c.font = PROMPT_FONT
    c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True, indent=1)
    c.border = THIN_BORDER
    ws.row_dimensions[row].height = 52
    end_col_num = ord(end_col) - ord("A") + 1
    # Style the answer cells BEFORE merging (merged cells become read-only)
    for col in range(2, end_col_num + 1):
        fc = ws.cell(row=row, column=col, value="")
        fc.fill = OFF_WHITE_FILL
        fc.border = THIN_BORDER
    if end_col_num > 1:
        ws.merge_cells(f"B{row}:{end_col}{row}")

# ─── Calendar data ──────────────────────────────────────────────────────────
WEEKS = [
    {
        "num": 1,
        "label": "WEEK 1 — Signal That Something Is Coming",
        "session": "Sunday, March 1, 2026",
        "posts": [
            {
                "date": date(2026, 3, 2), "day": "Monday",
                "account": "Personal Instagram", "series": "The Eye",
                "description": "Strongest Pinterest pin representing the brand's style philosophy — a single image that captures how you think about getting dressed.",
                "caption": "One sentence on why this image captures how you think about getting dressed.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 4), "day": "Wednesday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "Photo of your actual workspace this week — laptop, notebook, whatever is real. No staging.",
                "caption": "'Started documenting the build. More soon.'",
                "source": "Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 5), "day": "Thursday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Texture or material image from the Pinterest mood board — something representing the tactile quality Sylvan Hills is building toward.",
                "caption": "Minimal caption. Let the aesthetic speak.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 7), "day": "Saturday",
                "account": "Personal Instagram", "series": "The Eye",
                "description": "A menswear reference that influenced your thinking — a designer, era, or garment that shaped your point of view.",
                "caption": "What it taught you about restraint, quality, or intention.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 8), "day": "Sunday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Color palette or environmental image — muted, natural, quiet. Something that establishes the emotional register of the brand.",
                "caption": "No caption, or one-word mood at most.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
        ],
        "batch_tasks": [
            ("Plan week's content + pull Pinterest pins for 3 posts", "Both", "—", "30 min"),
            ("Shoot workspace photo for Wednesday Build post", "Personal", "The Build", "15 min"),
            ("Write captions for all 5 posts", "Both", "All series", "20 min"),
            ("Find/confirm Pinterest pins (Mon, Thu, Sat, Sun posts)", "Both", "Eye / World", "15 min"),
            ("Edit, format, and schedule Monday post", "Personal", "The Eye", "10 min"),
            ("Edit, format, and schedule Wednesday post", "Personal", "The Build", "10 min"),
            ("Edit, format, and schedule Thursday post", "@SylvanHillsofficial", "The World", "10 min"),
            ("Edit, format, and schedule Saturday post", "Personal", "The Eye", "10 min"),
            ("Edit, format, and schedule Sunday post", "@SylvanHillsofficial", "The World", "10 min"),
        ]
    },
    {
        "num": 2,
        "label": "WEEK 2 — Deepen the Aesthetic",
        "session": "Sunday, March 8, 2026",
        "posts": [
            {
                "date": date(2026, 3, 9), "day": "Monday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Outdoor or environmental image representing the Sylvan Hills man's world — grounded, quiet, intentional.",
                "caption": "Minimal. Let the world speak.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 10), "day": "Tuesday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "Something specific you worked on this week — a decision made, a resource found, a question you're wrestling with. Be specific and honest.",
                "caption": "Raw and honest. No polish. What actually happened this week.",
                "source": "Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 11), "day": "Wednesday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Close-up detail image — texture, craft, material. Training the audience's eye to what Sylvan Hills sees, before product exists.",
                "caption": "Short, precise observation about craft or material quality.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 13), "day": "Friday",
                "account": "Personal Instagram", "series": "The Eye",
                "description": "Non-fashion reference — architecture, interiors, or art — that shares Sylvan Hills DNA. Something that reveals how you see.",
                "caption": "Connect the visual reference to your design thinking or philosophy.",
                "source": "Pinterest / Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 15), "day": "Sunday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "Honest weekly reflection — what moved forward, what didn't. The most important post of the first two weeks.",
                "caption": "Candid reflection. Expected to be the most engaged post of the month.",
                "source": "Original", "status": "Not Started", "notes": "HIGH PRIORITY"
            },
        ],
        "batch_tasks": [
            ("Plan week's content + review Week 1 engagement", "Both", "—", "30 min"),
            ("Write Tuesday Build post (something specific worked on)", "Personal", "The Build", "20 min"),
            ("Write Sunday reflection post — most important post this week", "Personal", "The Build", "30 min"),
            ("Find Pinterest pins (Mon, Wed, Fri posts)", "Both", "World / Eye", "15 min"),
            ("Write captions for all 5 posts", "Both", "All series", "15 min"),
            ("Edit, format, and schedule all 5 posts", "Both", "All series", "40 min"),
        ]
    },
    {
        "num": 3,
        "label": "WEEK 3 — Name the Brand",
        "session": "Sunday, March 15, 2026",
        "posts": [
            {
                "date": date(2026, 3, 16), "day": "Monday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Lifestyle image — the brand's man made visual, without showing a face. Clothing, environment, posture. Quiet identity.",
                "caption": "One line, if any. The image does the work.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 17), "day": "Tuesday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "Name Sylvan Hills for the first time publicly. Origin of the name, what it means, why now. 4–5 sentences. Update bio to add @SylvanHillsofficial and link.",
                "caption": "Origin of name, meaning, why now. Update bio this same day.",
                "source": "Original", "status": "Not Started", "notes": "UPDATE BIO — add @SylvanHillsofficial to personal account bio"
            },
            {
                "date": date(2026, 3, 18), "day": "Wednesday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Mirror the name announcement on the official account. 'This is Sylvan Hills.' One line, clean image — original or strongest environmental pin.",
                "caption": "'This is Sylvan Hills.' One line, nothing more.",
                "source": "Original / Pinterest", "status": "Not Started", "notes": "Coordinate with Tuesday personal post"
            },
            {
                "date": date(2026, 3, 19), "day": "Thursday",
                "account": "Personal Instagram", "series": "The Eye",
                "description": "Design reference — furniture, architecture, or object — that shares Sylvan Hills DNA. Something connecting craft, restraint, and intention.",
                "caption": "Connect the visual reference to brand design philosophy.",
                "source": "Pinterest / Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 21), "day": "Saturday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Brand summary statement as a typography post: 'A place you go to come home to yourself.' Minimal design. Clean background, refined type.",
                "caption": "No caption needed. The statement is the caption.",
                "source": "Original", "status": "Not Started", "notes": "Design asset needed — create in Canva"
            },
        ],
        "batch_tasks": [
            ("Draft 'Name the Brand' post — Tuesday, most important this week", "Personal", "The Build", "45 min"),
            ("Draft 'This is Sylvan Hills.' announcement for official account (Wed)", "@SylvanHillsofficial", "The World", "15 min"),
            ("Design typography post: 'A place you go...' (Sat)", "@SylvanHillsofficial", "The World", "30 min"),
            ("Find Pinterest pins (Mon, Thu posts)", "Both", "World / Eye", "15 min"),
            ("Write all captions; update personal bio on Tuesday post day", "Personal", "The Build", "15 min"),
            ("Edit, format, and schedule all 5 posts", "Both", "All series", "30 min"),
        ]
    },
    {
        "num": 4,
        "label": "WEEK 4 — Close the Loop and Build Forward",
        "session": "Sunday, March 22, 2026",
        "posts": [
            {
                "date": date(2026, 3, 23), "day": "Monday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "What you're specifically working on this week — concrete and real, even if it's just research and decisions. Show the actual process.",
                "caption": "Concrete update. No vagueness. Name what you're actually doing.",
                "source": "Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 24), "day": "Tuesday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Material or fabric story sourced from Pinterest — what Sylvan Hills will be made from and the feeling those materials create. Natural, tactile, enduring.",
                "caption": "Material language. What you're building toward and why it matters.",
                "source": "Pinterest", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 25), "day": "Wednesday",
                "account": "Personal Instagram", "series": "The Eye",
                "description": "Style observation from personal life this week — something worn, noticed, or decided that confirmed your design direction.",
                "caption": "Personal, observational, honest. What you noticed and what it confirmed.",
                "source": "Original", "status": "Not Started", "notes": ""
            },
            {
                "date": date(2026, 3, 26), "day": "Thursday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Brand philosophy post — one core Sylvan Hills value expressed as a single clean sentence over a minimal image.",
                "caption": "One sentence. One value. Clean image. No explanation needed.",
                "source": "Original", "status": "Not Started", "notes": "Design asset needed"
            },
            {
                "date": date(2026, 3, 27), "day": "Friday",
                "account": "Personal Instagram", "series": "The Build",
                "description": "30-day reflection — where it started, what happened, what surprised you. The most important post of the month. Full transparency.",
                "caption": "Most important post of Month 1. Honest, complete, and forward-looking.",
                "source": "Original", "status": "Not Started", "notes": "HIGH PRIORITY — draft in advance"
            },
            {
                "date": date(2026, 3, 28), "day": "Saturday",
                "account": "@SylvanHillsofficial", "series": "The World",
                "description": "Four images from the month as a carousel — the brand world you've built, made visible as a whole. Curation over novelty.",
                "caption": "The world of Sylvan Hills, made visible. Four images. No words needed.",
                "source": "Pinterest", "status": "Not Started", "notes": "Select 4 strongest images from Month 1"
            },
        ],
        "batch_tasks": [
            ("Draft 30-day reflection post — most important post of the month (Fri)", "Personal", "The Build", "45 min"),
            ("Design brand philosophy post — one sentence, clean image (Thu)", "@SylvanHillsofficial", "The World", "20 min"),
            ("Select 4 strongest images for month-close carousel (Sat)", "@SylvanHillsofficial", "The World", "20 min"),
            ("Write Monday Build update — concrete, specific", "Personal", "The Build", "15 min"),
            ("Find Pinterest pins (Tue, Wed material/Eye posts)", "Both", "World / Eye", "15 min"),
            ("Edit, format, and schedule all 6 posts", "Both", "All series", "45 min"),
            ("Complete Monthly Reflection sheet", "—", "—", "30 min"),
        ]
    },
]

# ============================================================
#  XLSX WORKBOOK
# ============================================================
wb = openpyxl.Workbook()

# ── Sheet 1: Content Calendar ────────────────────────────────
ws1 = wb.active
ws1.title = "Content Calendar"

# Title rows
ws1.merge_cells("A1:I2")
t = ws1["A1"]
t.value = "SYLVAN HILLS — 30-DAY CONTENT CALENDAR"
t.font = TITLE_FONT
t.alignment = Alignment(horizontal="center", vertical="center")
t.fill = NORMAL_FILL
ws1.row_dimensions[1].height = 22
ws1.row_dimensions[2].height = 8

ws1.merge_cells("A3:I3")
s = ws1["A3"]
s.value = (
    "A place you go to come home to yourself.  ·  March 2026  ·  "
    "Batch session every Sunday evening  ·  ~7–9 posts / week across both accounts"
)
s.font = SUBTITLE_FONT
s.alignment = Alignment(horizontal="center", vertical="center")
s.fill = ALT_FILL
ws1.row_dimensions[3].height = 18

# Header row 4
COLS = ["Date", "Day of Week", "Account", "Series", "Content Description",
        "Caption Notes", "Content Source", "Status", "Notes"]
COL_W = [13, 12, 22, 14, 48, 44, 18, 14, 32]
for ci, (col_name, w) in enumerate(zip(COLS, COL_W), 1):
    c = ws1.cell(row=4, column=ci, value=col_name)
    hdr(c)
    ws1.column_dimensions[get_column_letter(ci)].width = w
ws1.row_dimensions[4].height = 30

ws1.freeze_panes = "A5"

row = 5
idx = 0
for week in WEEKS:
    # Week banner
    ws1.merge_cells(f"A{row}:I{row}")
    wc = ws1.cell(row=row, column=1, value=f"  {week['label']}")
    wc.fill = WEEK_FILL
    wc.font = WEEK_FONT
    wc.alignment = Alignment(horizontal="left", vertical="center")
    wc.border = THIN_BORDER
    ws1.row_dimensions[row].height = 22
    row += 1
    idx = 0

    for post in week["posts"]:
        alt = idx % 2 == 1
        vals = [
            post["date"].strftime("%b %d, %Y"),
            post["day"],
            post["account"],
            post["series"],
            post["description"],
            post["caption"],
            post["source"],
            post["status"],
            post["notes"],
        ]
        for ci, val in enumerate(vals, 1):
            c = ws1.cell(row=row, column=ci, value=val)
            body(c, alt=alt)
        ws1.row_dimensions[row].height = 65
        row += 1
        idx += 1

# Conditional formatting — Status column H
sr = f"H5:H{row}"
def cf_rule(val, fg, font_color=WHITE, bold=True):
    return CellIsRule(
        operator="equal",
        formula=[f'"{val}"'],
        fill=PatternFill("solid", fgColor=fg),
        font=Font(color=font_color, bold=bold),
    )
ws1.conditional_formatting.add(sr, cf_rule("Not Started", LIGHT_GREY, DARK_TEXT, False))
ws1.conditional_formatting.add(sr, cf_rule("In Progress", WARM_TAUPE))
ws1.conditional_formatting.add(sr, cf_rule("Scheduled",   SOFT_NAVY))
ws1.conditional_formatting.add(sr, cf_rule("Published",   SAGE_GREEN))


# ── Sheet 2: Weekly Batch Planner ───────────────────────────
ws2 = wb.create_sheet("Weekly Batch Planner")

ws2.merge_cells("A1:F2")
t2 = ws2["A1"]
t2.value = "SYLVAN HILLS — WEEKLY BATCH PLANNER"
t2.font = TITLE_FONT
t2.alignment = Alignment(horizontal="center", vertical="center")
t2.fill = NORMAL_FILL
ws2.row_dimensions[1].height = 22
ws2.row_dimensions[2].height = 8

ws2.merge_cells("A3:F3")
s2 = ws2["A3"]
s2.value = "Sunday evening · 2.5-hour session · Plan 30 min · Create 90 min · Edit & Schedule 30–60 min"
s2.font = SUBTITLE_FONT
s2.alignment = Alignment(horizontal="center", vertical="center")
s2.fill = ALT_FILL
ws2.row_dimensions[3].height = 18

batch_hdrs = ["✓", "Task / Post", "Account", "Series", "Est. Time", "Notes / Collected"]
batch_widths = [5, 42, 22, 16, 13, 40]
for ci, (h, w) in enumerate(zip(batch_hdrs, batch_widths), 1):
    c = ws2.cell(row=4, column=ci, value=h)
    hdr(c)
    ws2.column_dimensions[get_column_letter(ci)].width = w
ws2.row_dimensions[4].height = 28
ws2.freeze_panes = "A5"

row2 = 5
for week in WEEKS:
    # Week banner
    ws2.merge_cells(f"A{row2}:F{row2}")
    wc = ws2.cell(row=row2, column=1,
                  value=f"  WEEK {week['num']} — Batch Session: {week['session']}")
    wc.fill = HEADER_FILL
    wc.font = Font(name="Calibri", bold=True, color=WHITE, size=11)
    wc.alignment = Alignment(horizontal="left", vertical="center")
    wc.border = THIN_BORDER
    ws2.row_dimensions[row2].height = 24
    row2 += 1

    for ti, (task, acct, series, est) in enumerate(week["batch_tasks"]):
        alt = ti % 2 == 1
        vals = ["☐", task, acct, series, est, ""]
        for ci, val in enumerate(vals, 1):
            c = ws2.cell(row=row2, column=ci, value=val)
            body(c, alt=alt)
            if ci == 1:
                c.alignment = Alignment(horizontal="center", vertical="center")
        ws2.row_dimensions[row2].height = 22
        row2 += 1

    # Total row
    ws2.merge_cells(f"A{row2}:D{row2}")
    tc = ws2.cell(row=row2, column=1, value=f"  Week {week['num']} total")
    tc.fill = ALT_FILL
    tc.font = Font(name="Calibri", bold=True, color=DEEP_NAVY, size=10)
    tc.alignment = Alignment(horizontal="left", vertical="center", indent=1)
    tc.border = THIN_BORDER
    ec = ws2.cell(row=row2, column=5, value="~2.5 hrs")
    ec.fill = ALT_FILL
    ec.font = Font(name="Calibri", bold=True, color=DEEP_NAVY, size=10)
    ec.alignment = Alignment(horizontal="center", vertical="center")
    ec.border = THIN_BORDER
    nc = ws2.cell(row=row2, column=6, value="")
    nc.fill = ALT_FILL
    nc.border = THIN_BORDER
    ws2.row_dimensions[row2].height = 20
    row2 += 2


# ── Sheet 3: Content Library ─────────────────────────────────
ws3 = wb.create_sheet("Content Library")

ws3.merge_cells("A1:E2")
t3 = ws3["A1"]
t3.value = "SYLVAN HILLS — CONTENT LIBRARY"
t3.font = TITLE_FONT
t3.alignment = Alignment(horizontal="center", vertical="center")
t3.fill = NORMAL_FILL
ws3.row_dimensions[1].height = 22
ws3.row_dimensions[2].height = 8

ws3.merge_cells("A3:E3")
s3 = ws3["A3"]
s3.value = "Log every image, original shot, or curated asset here before assigning it to the calendar."
s3.font = SUBTITLE_FONT
s3.alignment = Alignment(horizontal="center", vertical="center")
s3.fill = ALT_FILL
ws3.row_dimensions[3].height = 18

lib_hdrs = ["Image Description", "Source", "Assigned To (Post Date & Series)", "Status", "Notes"]
lib_widths = [48, 20, 38, 16, 40]
for ci, (h, w) in enumerate(zip(lib_hdrs, lib_widths), 1):
    c = ws3.cell(row=4, column=ci, value=h)
    hdr(c)
    ws3.column_dimensions[get_column_letter(ci)].width = w
ws3.row_dimensions[4].height = 28
ws3.freeze_panes = "A5"

lib_rows = [
    ("Pinterest pin — single image capturing style philosophy", "Pinterest",
     "Mar 2 — The Eye (Wk 1 Mon)", "Reserved", "Pull from main mood board"),
    ("Workspace photo — laptop, notebook, current desk setup", "Original",
     "Mar 4 — The Build (Wk 1 Wed)", "Available", "Shoot during the week"),
    ("Texture / material close-up from Pinterest board", "Pinterest",
     "Mar 5 — The World (Wk 1 Thu)", "Reserved", "Select from 65-pin mood board"),
    ("Menswear reference — designer or garment that influenced thinking", "Pinterest",
     "Mar 7 — The Eye (Wk 1 Sat)", "Available", "Describe what it taught you"),
    ("Environmental color palette image — muted, natural tones", "Pinterest",
     "Mar 8 — The World (Wk 1 Sun)", "Available", "Focus on color over subject matter"),
    ("Outdoor / environmental lifestyle image — grounded, quiet", "Pinterest",
     "Mar 9 — The World (Wk 2 Mon)", "Available", "From existing mood board"),
    ("Non-fashion reference — architecture, interiors, or art", "Pinterest / Original",
     "Mar 13 — The Eye (Wk 2 Fri)", "Available", "Connect to brand design thinking"),
    ("Typography asset: 'A place you go to come home to yourself.'", "Original",
     "Mar 21 — The World (Wk 3 Sat)", "Not Started", "Create in Canva — minimal design"),
    ("Brand philosophy one-sentence post — design asset", "Original",
     "Mar 26 — The World (Wk 4 Thu)", "Not Started", "Design asset needed"),
    ("Month 1 carousel — 4 strongest images from March", "Pinterest",
     "Mar 28 — The World (Wk 4 Sat)", "Not Started", "Curate at end of month"),
]

for ri, row_data in enumerate(lib_rows):
    alt = ri % 2 == 1
    for ci, val in enumerate(row_data, 1):
        c = ws3.cell(row=5 + ri, column=ci, value=val)
        body(c, alt=alt)
    ws3.row_dimensions[5 + ri].height = 36

lib_status_range = f"D5:D{5+len(lib_rows)}"
ws3.conditional_formatting.add(lib_status_range, cf_rule("Available",  SAGE_GREEN))
ws3.conditional_formatting.add(lib_status_range, cf_rule("Reserved",   SOFT_NAVY))
ws3.conditional_formatting.add(lib_status_range, cf_rule("Used",       WARM_TAUPE))
ws3.conditional_formatting.add(lib_status_range, cf_rule("Not Started", LIGHT_GREY, DARK_TEXT, False))


# ── Sheet 4: Monthly Reflection ──────────────────────────────
ws4 = wb.create_sheet("Monthly Reflection")

ws4.column_dimensions["A"].width = 42
ws4.column_dimensions["B"].width = 20
ws4.column_dimensions["C"].width = 20
ws4.column_dimensions["D"].width = 20

ws4.merge_cells("A1:D2")
t4 = ws4["A1"]
t4.value = "SYLVAN HILLS — MONTH 1 REFLECTION"
t4.font = TITLE_FONT
t4.alignment = Alignment(horizontal="center", vertical="center")
t4.fill = NORMAL_FILL
ws4.row_dimensions[1].height = 22
ws4.row_dimensions[2].height = 8

ws4.merge_cells("A3:D3")
s4 = ws4["A3"]
s4.value = "Complete at the end of each week. Use month-close to inform Month 2 strategy and cadence."
s4.font = SUBTITLE_FONT
s4.alignment = Alignment(horizontal="center", vertical="center")
s4.fill = ALT_FILL
ws4.row_dimensions[3].height = 18
ws4.freeze_panes = "A5"

r4 = 5
week_ranges = ["March 2–8", "March 9–15", "March 16–22", "March 23–28"]
week_prompts = [
    "What went well this week?",
    "What stalled or didn't happen?",
    "One thing to do differently next week.",
    "Confidence rating this week (1–10):",
]

for wn, wr in enumerate(week_ranges, 1):
    section_hdr(ws4, r4, f"  WEEK {wn} REFLECTION — {wr}")
    r4 += 1
    for pi, p in enumerate(week_prompts):
        prompt_row(ws4, r4, p, alt=(pi % 2 == 1))
        r4 += 1
    r4 += 1

section_hdr(ws4, r4, "  MONTH 1 SUCCESS METRICS")
r4 += 1
metrics = [
    "Total followers gained — Personal Instagram:",
    "Total followers gained — @SylvanHillsofficial:",
    "Most engaged post of the month:",
    "Least engaged post and why:",
    "Content series that felt most natural to create:",
    "Content series that felt most difficult:",
    "Overall consistency rating (1–10):",
]
for mi, m in enumerate(metrics):
    prompt_row(ws4, r4, m, alt=(mi % 2 == 1))
    r4 += 1

r4 += 1
section_hdr(ws4, r4, "  MONTH 2 STRATEGIC ADJUSTMENTS")
r4 += 1
for sub_r in range(r4, r4 + 3):
    if sub_r == r4:
        lc = ws4.cell(row=sub_r, column=1, value="Based on Month 1 learnings, what changes for Month 2?")
        lc.fill = NORMAL_FILL
        lc.font = PROMPT_FONT
        lc.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True, indent=1)
        lc.border = THIN_BORDER
    else:
        lc = ws4.cell(row=sub_r, column=1, value="")
        lc.fill = NORMAL_FILL
        lc.border = THIN_BORDER
    ws4.row_dimensions[sub_r].height = 55
    # Style answer cells BEFORE merging
    for col in range(2, 5):
        fc = ws4.cell(row=sub_r, column=col, value="")
        fc.fill = OFF_WHITE_FILL
        fc.border = THIN_BORDER
    ws4.merge_cells(f"B{sub_r}:D{sub_r}")

# Save xlsx
XLSX_PATH = "/home/user/SH2026/content/sylvan-hills-content-calendar.xlsx"
wb.save(XLSX_PATH)
print(f"Saved: {XLSX_PATH}")


# ============================================================
#  WORD DOCUMENT
# ============================================================
doc = Document()

# ── Page margins ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

def set_para_spacing(para, before=0, after=6, line_rule=None, line=None):
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), line_rule or "auto")
    pPr.append(spacing)

def rgb(hex_str):
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)

def add_rule(doc, color_hex=TAUPE_BORDER):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)
    set_para_spacing(p, before=60, after=60)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color=TAUPE_BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

# ── Cover page ───────────────────────────────────────────────
# Vertical spacing to center
for _ in range(6):
    sp = doc.add_paragraph()
    set_para_spacing(sp, before=0, after=0)

brand_p = doc.add_paragraph()
brand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = brand_p.add_run("SYLVAN HILLS")
run.font.name = "Calibri"
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = rgb(DEEP_NAVY)
set_para_spacing(brand_p, before=0, after=24)

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = tag_p.add_run("A place you go to come home to yourself.")
run2.font.name = "Calibri"
run2.font.size = Pt(14)
run2.font.italic = True
run2.font.color.rgb = rgb(MID_TAUPE)
set_para_spacing(tag_p, before=0, after=48)

add_rule(doc)

doc_title_p = doc.add_paragraph()
doc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = doc_title_p.add_run("30-Day Content Calendar")
run3.font.name = "Calibri"
run3.font.size = Pt(18)
run3.font.bold = True
run3.font.color.rgb = rgb(DEEP_NAVY)
set_para_spacing(doc_title_p, before=48, after=12)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = date_p.add_run("March 2026  ·  Prepared February 24, 2026")
run4.font.name = "Calibri"
run4.font.size = Pt(11)
run4.font.color.rgb = rgb(MID_TAUPE)
set_para_spacing(date_p, before=0, after=0)

doc.add_page_break()

# ── Brand Context ────────────────────────────────────────────
def doc_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = rgb(DEEP_NAVY)
    set_para_spacing(p, before=120, after=60)
    return p

def doc_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = rgb(SOFT_NAVY)
    set_para_spacing(p, before=80, after=30)
    return p

def doc_body(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.italic = italic
    run.font.color.rgb = rgb(DARK_TEXT)
    set_para_spacing(p, before=0, after=36)
    return p

def doc_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(DARK_TEXT)
    set_para_spacing(p, before=0, after=18)
    return p

doc_h1(doc, "Brand Context")

doc_body(doc,
    "Sylvan Hills is a founder-led menswear brand built around the positioning: "
    "'A place you go to come home to yourself.' The brand targets men who dress from "
    "the inside out — not trend followers, but men with their own taste and quiet "
    "self-assurance. The founder is building the brand part-time while employed full-time, "
    "with a planned transition to full-time in June 2026."
)

doc_h2(doc, "The Two Accounts")
doc_bullet(doc, "Personal Instagram (850 followers) — founder's personal account, handle TBC")
doc_bullet(doc, "@SylvanHillsofficial — Sylvan Hills brand account")

doc_h2(doc, "The Three Content Series")
doc_bullet(doc,
    "The Build (Personal) — Behind-the-scenes founder documentation. "
    "Raw, honest, process-driven. Shows the real journey of building Sylvan Hills in real time."
)
doc_bullet(doc,
    "The Eye (Personal) — Taste curation and point of view. Images, references, and "
    "observations that establish the founder's credibility as a tastemaker. "
    "Curated Pinterest images with original captions."
)
doc_bullet(doc,
    "The World (@SylvanHillsofficial) — Brand aesthetic and emotional territory. "
    "Curated images establishing the visual language of the brand before product exists. "
    "Sourced primarily from the existing 65-pin Pinterest mood board."
)

doc_h2(doc, "Posting Cadence")
doc_bullet(doc, "Personal / The Build: 2–3 posts per week")
doc_bullet(doc, "Personal / The Eye: 2 posts per week")
doc_bullet(doc, "@SylvanHillsofficial / The World: 3–4 posts per week")
doc_bullet(doc, "Total: ~7–9 posts per week across both accounts")

doc_body(doc,
    "Content is batched in a single 2.5-hour session every Sunday evening. "
    "This structure allows for intentional creation rather than reactive posting — "
    "consistent with the brand's philosophy of doing fewer things, better.",
    italic=True
)

doc_h2(doc, "Strategic Intention")
doc_body(doc,
    "Month 1 is not about reach. It is about establishing signal, building discipline, "
    "and beginning to shape the aesthetic and voice of Sylvan Hills before product exists. "
    "The 30-day plan moves through four phases: signaling that something is coming, "
    "deepening the aesthetic, naming the brand publicly, and closing the first chapter "
    "while setting up Month 2."
)

add_rule(doc)

# ── Week-by-week calendar ─────────────────────────────────────
doc_h1(doc, "30-Day Content Calendar")

SERIES_COLORS = {
    "The Build": WARM_TAUPE,
    "The Eye":   SOFT_NAVY,
    "The World": "4A7060",
}

for week in WEEKS:
    doc_h2(doc, week["label"])

    for post in week["posts"]:
        # Post header table (1 row, all metadata)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        widths = [Inches(1.1), Inches(0.9), Inches(1.4), Inches(1.0), Inches(1.5)]
        row_cells = tbl.rows[0].cells
        headers_data = [
            post["date"].strftime("%b %d"),
            post["day"],
            post["account"],
            post["series"],
            post["source"],
        ]
        header_bgs = [DEEP_NAVY, DEEP_NAVY, DEEP_NAVY, SERIES_COLORS.get(post["series"], DEEP_NAVY), DEEP_NAVY]

        for ci, (cell, val, bg, w) in enumerate(zip(row_cells, headers_data, header_bgs, widths)):
            cell.width = w
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = rgb(WHITE)
            p_pPr = p._p.get_or_add_pPr()
            sp_el = OxmlElement("w:spacing")
            sp_el.set(qn("w:before"), "30")
            sp_el.set(qn("w:after"), "30")
            p_pPr.append(sp_el)

        # Content row (description + caption)
        tbl2 = doc.add_table(rows=2, cols=2)
        tbl2.style = "Table Grid"
        tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

        field_labels = ["Content", "Caption Notes"]
        field_values = [post["description"], post["caption"]]
        field_bg     = [SOFT_NEUTRAL, CREAM]

        for ri2, (label, value, bg) in enumerate(zip(field_labels, field_values, field_bg)):
            label_cell = tbl2.rows[ri2].cells[0]
            value_cell = tbl2.rows[ri2].cells[1]

            label_cell.width = Inches(1.1)
            value_cell.width = Inches(4.9)

            set_cell_bg(label_cell, bg)
            set_cell_bg(value_cell, bg)
            set_cell_borders(label_cell)
            set_cell_borders(value_cell)

            lp = label_cell.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.name = "Calibri"
            lr.font.bold = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = rgb(DEEP_NAVY)
            lp_pPr = lp._p.get_or_add_pPr()
            lsp = OxmlElement("w:spacing")
            lsp.set(qn("w:before"), "30")
            lsp.set(qn("w:after"), "30")
            lp_pPr.append(lsp)

            vp = value_cell.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.name = "Calibri"
            vr.font.size = Pt(9.5)
            vr.font.color.rgb = rgb(DARK_TEXT)
            if ri2 == 1:
                vr.font.italic = True
                vr.font.color.rgb = rgb(MID_TAUPE)
            vp_pPr = vp._p.get_or_add_pPr()
            vsp = OxmlElement("w:spacing")
            vsp.set(qn("w:before"), "30")
            vsp.set(qn("w:after"), "30")
            vp_pPr.append(vsp)

        # Notes row if any
        if post["notes"]:
            tbl3 = doc.add_table(rows=1, cols=2)
            tbl3.style = "Table Grid"
            nc1 = tbl3.rows[0].cells[0]
            nc2 = tbl3.rows[0].cells[1]
            nc1.width = Inches(1.1)
            nc2.width = Inches(4.9)
            note_bg = "FFF0C4"
            set_cell_bg(nc1, note_bg)
            set_cell_bg(nc2, note_bg)
            set_cell_borders(nc1)
            set_cell_borders(nc2)
            np1 = nc1.paragraphs[0]
            nr1 = np1.add_run("Note")
            nr1.font.name = "Calibri"
            nr1.font.bold = True
            nr1.font.size = Pt(9)
            nr1.font.color.rgb = rgb("7A5F00")
            np1_pPr = np1._p.get_or_add_pPr()
            ns1 = OxmlElement("w:spacing")
            ns1.set(qn("w:before"), "30")
            ns1.set(qn("w:after"), "30")
            np1_pPr.append(ns1)
            np2 = nc2.paragraphs[0]
            nr2 = np2.add_run(post["notes"])
            nr2.font.name = "Calibri"
            nr2.font.size = Pt(9)
            nr2.font.bold = True
            nr2.font.color.rgb = rgb("7A5F00")
            np2_pPr = np2._p.get_or_add_pPr()
            ns2 = OxmlElement("w:spacing")
            ns2.set(qn("w:before"), "30")
            ns2.set(qn("w:after"), "30")
            np2_pPr.append(ns2)

        sp_p = doc.add_paragraph()
        set_para_spacing(sp_p, before=0, after=60)

    add_rule(doc)

# ── Monthly Reflection ────────────────────────────────────────
doc.add_page_break()
doc_h1(doc, "Monthly Reflection Template")

doc_body(doc,
    "Complete each week section on the Sunday batch session. Complete the Month 1 metrics "
    "section at the end of March before planning Month 2. Be honest — this document is for "
    "you, not for an audience."
)

for wn, wr in enumerate(week_ranges, 1):
    doc_h2(doc, f"Week {wn} Reflection — {wr}")
    for prompt in week_prompts:
        tbl_r = doc.add_table(rows=1, cols=2)
        tbl_r.style = "Table Grid"
        lc = tbl_r.rows[0].cells[0]
        vc = tbl_r.rows[0].cells[1]
        lc.width = Inches(2.4)
        vc.width = Inches(3.6)
        set_cell_bg(lc, CREAM)
        set_cell_bg(vc, OFF_WHITE)
        set_cell_borders(lc)
        set_cell_borders(vc)
        lp = lc.paragraphs[0]
        lr = lp.add_run(prompt)
        lr.font.name = "Calibri"
        lr.font.bold = True
        lr.font.size = Pt(9.5)
        lr.font.color.rgb = rgb(DEEP_NAVY)
        lp_pPr = lp._p.get_or_add_pPr()
        lsp = OxmlElement("w:spacing")
        lsp.set(qn("w:before"), "60")
        lsp.set(qn("w:after"), "60")
        lp_pPr.append(lsp)
        vp = vc.paragraphs[0]
        vp.add_run("")
        vp_pPr = vp._p.get_or_add_pPr()
        vsp = OxmlElement("w:spacing")
        vsp.set(qn("w:before"), "60")
        vsp.set(qn("w:after"), "60")
        vp_pPr.append(vsp)
        sp_p2 = doc.add_paragraph()
        set_para_spacing(sp_p2, before=0, after=18)

doc_h2(doc, "Month 1 Success Metrics")
for metric in metrics:
    tbl_m = doc.add_table(rows=1, cols=2)
    tbl_m.style = "Table Grid"
    mc1 = tbl_m.rows[0].cells[0]
    mc2 = tbl_m.rows[0].cells[1]
    mc1.width = Inches(2.8)
    mc2.width = Inches(3.2)
    set_cell_bg(mc1, SOFT_NEUTRAL)
    set_cell_bg(mc2, OFF_WHITE)
    set_cell_borders(mc1)
    set_cell_borders(mc2)
    mp1 = mc1.paragraphs[0]
    mr1 = mp1.add_run(metric)
    mr1.font.name = "Calibri"
    mr1.font.bold = True
    mr1.font.size = Pt(9.5)
    mr1.font.color.rgb = rgb(DEEP_NAVY)
    mp1_pPr = mp1._p.get_or_add_pPr()
    ms1 = OxmlElement("w:spacing")
    ms1.set(qn("w:before"), "50")
    ms1.set(qn("w:after"), "50")
    mp1_pPr.append(ms1)
    mp2 = mc2.paragraphs[0]
    mp2.add_run("")
    sp_m = doc.add_paragraph()
    set_para_spacing(sp_m, before=0, after=18)

doc_h2(doc, "Month 2 Strategic Adjustments")
doc_body(doc,
    "Based on Month 1 learnings — what changes in Month 2? What series gets more investment? "
    "What cadence proved sustainable? What surprised you about what resonated?",
    italic=True
)
adj_tbl = doc.add_table(rows=1, cols=1)
adj_tbl.style = "Table Grid"
adj_c = adj_tbl.rows[0].cells[0]
set_cell_bg(adj_c, OFF_WHITE)
set_cell_borders(adj_c)
adj_p = adj_c.paragraphs[0]
adj_p.add_run("")
adj_pPr = adj_p._p.get_or_add_pPr()
adj_sp = OxmlElement("w:spacing")
adj_sp.set(qn("w:before"), "0")
adj_sp.set(qn("w:after"), "0")
adj_sp.set(qn("w:line"), "1200")
adj_sp.set(qn("w:lineRule"), "exact")
adj_pPr.append(adj_sp)

# ── Footer note ──────────────────────────────────────────────
add_rule(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    "Sylvan Hills  ·  sylvanhills.com  ·  @SylvanHillsofficial  ·  "
    "Clothing for the man who dresses from the inside out."
)
fr.font.name = "Calibri"
fr.font.size = Pt(8)
fr.font.italic = True
fr.font.color.rgb = rgb(MID_TAUPE)
set_para_spacing(footer_p, before=30, after=0)

DOCX_PATH = "/home/user/SH2026/content/sylvan-hills-content-calendar.docx"
doc.save(DOCX_PATH)
print(f"Saved: {DOCX_PATH}")
print("Done.")
